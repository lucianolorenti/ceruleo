
import logging

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.table import Table
from pandas.api.types import is_numeric_dtype
from rul_pm.graphics.control_charts import plot_ewma_
from savona.exporter.docx_ import add_image_from_figure, table_from_pandas

logger = logging.getLogger(__name__)

plt.rcParams.update({'figure.max_open_warning': 0})


class Report:
    def generate(self):
        pass

    def save(self):
        pass


class FeatureReport(Report):
    def __init__(self, dataset, title='', output_file_name='features'):
        self.dataset = dataset
        self.df = dataset.toPandas()
        self.document = Document()
        self.title = title
        self._columns_type()
        self.footer()
        self.output_file_name = output_file_name

    def footer(self):
        section = self.document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.style = self.document.styles["Footer"]

    def _numerical_features_names(self):
        return [
            f for f in self.df.columns
            if is_numeric_dtype(self.df[f])]

    def _columns_type(self):
        self.columns_type = self.df.dtypes.reset_index()
        self.columns_type.columns = ['Feature', 'Type']

    def generate(self):
        self.document.add_heading(self.title, 0)
        self.feature_list()
        self.correlation()
        self._numerical_features()
        self.save()

    def feature_importance(self):
        pass

    def _features_info(self):
        df = self.df.isnull().mean().reset_index()
        df.columns = ['Feature', 'Null proportion']
        self.document.add_heading('Features information', 2)
        self.document.add_heading('Numerical features', 3)

        summary_numerical = (self
                             .df[self._numerical_features_names()]
                             .describe()
                             .T
                             .loc[:, ['std']]
                             .rename_axis('Feature')
                             .reset_index())

        df_numerical = (
            df
            .loc[df['Feature'].isin(self._numerical_features_names()), :]
            .merge(summary_numerical, on='Feature')
            .sort_values(by=['Null proportion', 'std']))

        table_from_pandas(self.document, df_numerical)

        self.document.add_heading('Categorical features', 3)
        categorical_features = (self
                                .df
                                .columns
                                .difference(
                                    self._numerical_features_names())
                                .values
                                .tolist())
        df_categorical = df[df['Feature'].isin(categorical_features)]

        summary_categorical = (self
                               .df
                               .loc[:, categorical_features]
                               .describe(include='all')
                               .T
                               .loc[:, ['unique', 'top', 'freq']]
                               .rename_axis('Feature')
                               .reset_index())
        df_categorical = (
            df_categorical
            .merge(summary_categorical, on='Feature')
            .sort_values(by=['Null proportion']))
        table_from_pandas(self.document, df_categorical)

    def ewma(self):
        pass

    def _numerical_features(self):
        bins = 15
        lambdas = [0.5, 1.0]
        colors = ['#9D8206', '#FC7A1E', '#B52C6E']
        self.document.add_heading('Numerical Features', 2)
        numerical_features = [
            f for f in self.df.columns
            if is_numeric_dtype(self.df[f])]
        for feature in numerical_features:
            not_null = self.df[feature].isnull().mean() < 0.5
            not_constant = self.df[feature].std() > 0
            if not(not_null and not_constant):
                continue
            self.document.add_heading(feature, 3)
            fig, ax = plt.subplots(1, 1, figsize=(15, 7))

            histogram = np.zeros((3, bins))
            for i, (l, color) in enumerate(zip(lambdas, colors)):
                LCL, mean, UCL, _ = plot_ewma_(
                    np.array(range(self.df.shape[0])),
                    self.df[feature].values,
                    feature,
                    ' ',
                    lambda_=l,
                    ax=ax,
                    color=color)

                for life in self.dataset:
                    p = (life[feature] <= LCL) | (life[feature] >= UCL)
                    b = np.linspace(0, life.shape[0], bins+1)
                    for j, bin_idx in enumerate(range(len(b)-1)):
                        histogram[i,
                                  j] += np.sum(p[int(b[bin_idx]): int(b[bin_idx+1])])

            add_image_from_figure(self.document, fig)

            for j, lamb in enumerate(lambdas):
                self.document.add_heading(
                    f'Proportion of samples outside the range lambda = {lamb}', 3)
                if np.sum(histogram[j, :]) > 0:
                    histogram[j, :] /= np.sum(histogram[j, :])
                else:
                    histogram[j, :] = 0
                b = np.linspace(0, 100, bins+1)
                labels = [
                    f'[{np.round(b[k], 2)}, {np.round(b[k+1], 2)})' for k in range(len(b)-1)]

                fig, ax = plt.subplots(1, 1, figsize=(15, 9))
                ax.bar(x=np.array(range(histogram.shape[1])),
                       height=histogram[j, :])

                ax.set_xlabel('Life percentage')
                ax.set_ylabel('Proportion of samples outside limits')
                ax.set_xticklabels(labels)
                ax.set_xticks(np.array(range(histogram.shape[1])))
                add_image_from_figure(self.document, fig)

    def feature_list(self):
        self.document.add_heading('Features', 1)
        self._features_info()

    def correlation(self):
        logger.info('Correlation')
        self.document.add_heading('Correlation', 1)
        corr_matrix = (self
                       .df[self._numerical_features_names()]
                       .corr()
                       .abs())

        sol = (corr_matrix
               .where(np.triu(np.ones(corr_matrix.shape), k=1).astype(np.bool))
               .stack()
               .sort_values(ascending=False)
               .rename_axis(['Feature 1', 'Feature 2'])
               .reset_index())
        sol.columns = ['Feature 1', 'Feature 2', 'Abs correlation']
        sol = sol[sol['Abs correlation'] > 0.8]
        sol['Abs correlation'] = sol['Abs correlation'].round(3)
        table_from_pandas(self.document, sol)

    def save(self):
        self.document.save(f'{self.output_file_name}.docx')

    def alarms(self):
        self.alarm_columns
